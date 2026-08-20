import io

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

from gerapop.constants import PASSO_COL_WIDTH_CM, ValidationMessage
from gerapop.models import PopData
from gerapop.services.documento import titulo_para_header
from gerapop.services.docx import gerar_docx
from gerapop.services.docx.styles import set_cell_shading


def test_set_cell_shading_nao_duplica() -> None:
    doc = Document()
    cell = doc.add_table(rows=1, cols=1).rows[0].cells[0]

    set_cell_shading(cell, "D9D9D9")
    set_cell_shading(cell, "D9D9D9")

    assert len(cell._tc.get_or_add_tcPr().findall(qn("w:shd"))) == 1


def test_gerar_docx_minimo(pop_minimo: PopData) -> None:
    content = gerar_docx(pop_minimo).getvalue()

    assert content[:2] == b"PK"
    assert len(content) > 1000


def test_validacao_campos_obrigatorios(pop_invalido: PopData) -> None:
    errors = pop_invalido.validate()

    assert ValidationMessage.NOME_OBRIGATORIO in errors
    assert ValidationMessage.OBJETIVO_OBRIGATORIO in errors
    assert ValidationMessage.CODIGO_OBRIGATORIO in errors
    assert ValidationMessage.AREA_OBRIGATORIA in errors


def test_output_filename(pop_minimo: PopData) -> None:
    pop_minimo.nome_pop = "Registro de Manobras no Sistema"

    assert pop_minimo.output_filename().startswith("POP-OPE-001_")
    assert pop_minimo.output_filename().endswith(".docx")


def _docx_estrutura(pop: PopData) -> tuple[list[str], list]:
    doc = Document(io.BytesIO(gerar_docx(pop).getvalue()))
    return [p.text for p in doc.paragraphs if p.text.strip()], doc.tables


def _texto_tabela(table) -> str:
    return table.rows[0].cells[0].text


def test_numercao_plana_e_aviso_no_escopo(pop_minimo: PopData) -> None:
    pop_minimo.aviso = "Este POP não contempla o anúncio do navio."
    pop_minimo.secoes = [
        {
            "titulo": "Acesso ao Módulo de Manobras",
            "responsavel": "",
            "passos": ["Abrir o sistema."],
            "campos": [],
        },
        {
            "titulo": "Procedimento – Atracação",
            "responsavel": "",
            "passos": ["Localizar o navio."],
            "campos": [],
        },
    ]

    texts, tables = _docx_estrutura(pop_minimo)

    assert "1.  Objetivo" in texts
    assert "Padronizar o registro de manobras." in texts
    assert "2.  Campo de Aplicação" in texts
    assert "Aplica-se à equipe de operações." in texts
    assert "3.  Definições" in texts
    assert "4.  Procedimento" in texts
    assert any("4.1" in t for t in texts)
    assert any("4.2" in t for t in texts)
    assert "5.  Regras e Restrições" in texts
    assert "6.  Consulta e Relatórios" in texts

    # Histórico de revisões e aprovação no topo
    assert any("Rev." in _texto_tabela(t) or "Revisão" in _texto_tabela(t) for t in tables)
    assert any("Elaborado por" in _texto_tabela(t) for t in tables)

    aviso_table = next(t for t in tables if "ATENÇÃO" in _texto_tabela(t))
    assert "ATENÇÃO: Este POP não contempla o anúncio do navio." in _texto_tabela(aviso_table)

    regra_table = next(t for t in tables if len(t.columns) == 2 and _texto_tabela(t) == "Regra")
    assert regra_table.rows[1].cells[1].text == "Não executar sem autorização."

    consulta_table = next(t for t in tables if len(t.columns) == 1 and "Menu" in _texto_tabela(t))
    assert consulta_table.rows[0].cells[0].text == "Menu > Operações > Manobras"


def test_aviso_vem_depois_do_escopo(pop_minimo: PopData) -> None:
    pop_minimo.aviso = "Procedimento condicionado."

    doc = Document(io.BytesIO(gerar_docx(pop_minimo).getvalue()))
    seq: list[str] = []
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
            if text:
                seq.append(text)
        elif child.tag == qn("w:tbl"):
            text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
            if text:
                seq.append(text)

    escopo = seq.index("Aplica-se à equipe de operações.")
    assert any("ATENÇÃO: Procedimento condicionado." in s for s in seq[escopo:])


def test_campos_obrigatorios_por_secao(pop_minimo: PopData) -> None:
    pop_minimo.secoes[0]["campos"] = [
        {"campo": "Berço", "descricao": "Número do berço de atracação designado."},
        {"campo": "Data e Hora", "descricao": "Data e hora efetiva da atracação."},
    ]

    texts, tables = _docx_estrutura(pop_minimo)

    assert any("Campos obrigatórios" in t for t in texts)
    campos_table = next(t for t in tables if _texto_tabela(t) == "Campo" and len(t.columns) == 2)
    assert [row.cells[0].text for row in campos_table.rows[1:]] == ["Berço", "Data e Hora"]
    assert campos_table.rows[1].cells[1].text == "Número do berço de atracação designado."


def test_secao_sem_campos_nao_gera_subtabela(pop_minimo: PopData) -> None:
    texts, tables = _docx_estrutura(pop_minimo)

    assert not any("Campos obrigatórios" in t for t in texts)
    assert not any(_texto_tabela(t) == "Campo" for t in tables)


def test_secao_antiga_sem_chave_campos(pop_minimo: PopData) -> None:
    pop_minimo.secoes[0] = {
        "titulo": "Atracação",
        "responsavel": "",
        "passos": ["Localizar o navio."],
        "campos": [],
    }

    texts, _ = _docx_estrutura(pop_minimo)

    assert "4.  Procedimento" in texts
    assert any("4.1" in t and "Atracação" in t for t in texts)


def _docx_completo(pop_minimo: PopData) -> Document:
    pop_minimo.aviso = "Procedimento condicionado à janela de maré."
    pop_minimo.definicoes = [
        {"termo": "TOS", "definicao": "Terminal Operating System"},
        {"termo": "SEV", "definicao": "Sistema Eletrônico de Vistorias"},
    ]
    pop_minimo.secoes[0]["campos"] = [
        {"campo": "Berço", "descricao": "Número do berço designado."},
        {"campo": "Data e Hora", "descricao": "Data e hora efetiva da atracação."},
    ]
    pop_minimo.regras = [
        "Não executar sem autorização.",
        "Usar EPI obrigatório.",
        "Comunicar qualquer anomalia ao supervisor.",
    ]
    return Document(io.BytesIO(gerar_docx(pop_minimo).getvalue()))


def _largura_util(doc: Document) -> int:
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def test_todas_tabelas_somam_a_largura_util(pop_minimo: PopData) -> None:
    doc = _docx_completo(pop_minimo)
    largura_util = _largura_util(doc)

    assert len(doc.tables) >= 6
    for table in doc.tables:
        larguras = [cell.width for cell in table.rows[0].cells]
        assert sum(larguras) == largura_util


def test_tabela_passos_numero_estreita_e_descricao_restante(pop_minimo: PopData) -> None:
    doc = _docx_completo(pop_minimo)
    largura_util = _largura_util(doc)

    passos_table = next(
        t for t in doc.tables if len(t.columns) == 2 and t.rows[0].cells[0].text == "#"
    )
    largura_numero = passos_table.rows[1].cells[0].width
    largura_descricao = passos_table.rows[1].cells[1].width
    assert abs(largura_numero - Cm(PASSO_COL_WIDTH_CM)) < 635
    assert abs(largura_descricao - (largura_util - Cm(PASSO_COL_WIDTH_CM))) < 635


def test_regras_numeradas_sequencialmente(pop_minimo: PopData) -> None:
    doc = _docx_completo(pop_minimo)

    regras_table = next(t for t in doc.tables if _texto_tabela(t) == "Regra")
    assert [row.cells[0].text for row in regras_table.rows[1:]] == ["R1", "R2", "R3"]
    assert regras_table.rows[3].cells[1].text == "Comunicar qualquer anomalia ao supervisor."


def test_passos_sub_cabecalho_e_resposta_sistema(pop_minimo: PopData) -> None:
    pop_minimo.secoes[0]["passos"] = [
        "Tela 6002 – Programação de Saída",
        "Abrir o sistema.",
        "Sistema exibe: 'Bem-vindo!' Clicar no botão 'Ok'.",
    ]

    _, tables = _docx_estrutura(pop_minimo)
    passos_table = next(t for t in tables if len(t.columns) == 2 and t.rows[0].cells[0].text == "#")

    sub_cell = passos_table.rows[1].cells[0]
    assert sub_cell.text == "Tela 6002 – Programação de Saída"
    assert sub_cell._tc.tcPr.find(qn("w:gridSpan")) is not None
    assert sub_cell.paragraphs[0].runs[0].bold

    sys_row = passos_table.rows[3]
    assert sys_row.cells[0].text == "—"
    assert all(run.italic for run in sys_row.cells[1].paragraphs[0].runs)


def test_aspas_simples_viram_negrito(pop_minimo: PopData) -> None:
    pop_minimo.secoes[0]["passos"] = ["Clicar no botão 'Novo' e depois em 'Gravar'."]

    _, tables = _docx_estrutura(pop_minimo)
    passos_table = next(t for t in tables if len(t.columns) == 2 and t.rows[0].cells[0].text == "#")

    runs = passos_table.rows[1].cells[1].paragraphs[0].runs
    assert [(run.text, run.bold) for run in runs] == [
        ("Clicar no botão ", False),
        ("Novo", True),
        (" e depois em ", False),
        ("Gravar", True),
        (".", False),
    ]


def test_larguras_validas_com_sub_cabecalho_na_primeira_linha(pop_minimo: PopData) -> None:
    pop_minimo.secoes[0]["passos"] = [
        "Tela 6002 – X",
        "Abrir.",
        "Sistema exibe: 'ok' Clicar em 'Sair'.",
    ]
    doc = Document(io.BytesIO(gerar_docx(pop_minimo).getvalue()))
    largura_util = _largura_util(doc)

    passos_table = next(
        t for t in doc.tables if len(t.columns) == 2 and t.rows[0].cells[0].text == "#"
    )
    primeira_normal = next(
        row
        for row in passos_table.rows
        if not any(
            (c._tc.tcPr is not None and c._tc.tcPr.find(qn("w:gridSpan")) is not None)
            for c in row.cells
        )
    )
    assert sum(cell.width for cell in primeira_normal.cells) == largura_util
    for row in passos_table.rows:
        for cell in row.cells:
            if cell._tc.tcPr is not None and cell._tc.tcPr.find(qn("w:gridSpan")) is not None:
                continue
            assert cell.width < largura_util


def test_rodape_com_codigo_versao_e_campo_pagina(pop_minimo: PopData) -> None:
    doc = Document(io.BytesIO(gerar_docx(pop_minimo).getvalue()))
    paragraph = doc.sections[0].footer.paragraphs[0]

    texto = "".join(node.text or "" for node in paragraph._p.iter(qn("w:t")))
    assert "POP-OPE-001" in texto
    assert "Registro de Manobras" in texto
    assert "Versão 01" in texto

    campos = paragraph._p.findall(qn("w:fldSimple"))
    assert campos and campos[0].get(qn("w:instr")) == "PAGE"


def test_matriz_fluxo_etapa_registro_atividade(pop_minimo: PopData) -> None:
    pop_minimo.matriz_responsabilidades = [
        {
            "etapa": "1",
            "registro": "Ficha de oportunidade",
            "atividade": "Criar e registrar cliente",
            "responsavel": "Comercial",
        },
        {
            "etapa": "2",
            "registro": "Avaliação comercial",
            "atividade": "Avaliar viabilidade",
            "responsavel": "Comitê comercial",
        },
    ]

    _, tables = _docx_estrutura(pop_minimo)
    matriz = next(t for t in tables if len(t.columns) == 4 and _texto_tabela(t) == "Etapa")
    assert [c.text for c in matriz.rows[0].cells] == [
        "Etapa",
        "Registro",
        "Atividade",
        "Responsável",
    ]
    assert matriz.rows[1].cells[1].text == "Ficha de oportunidade"
    assert matriz.rows[2].cells[3].text == "Comitê comercial"


def test_passos_com_responsavel_por_linha(pop_minimo: PopData) -> None:
    pop_minimo.secoes = [
        {
            "titulo": "Prospecção",
            "responsaveis": ["Comercial", "Logística"],
            "passos": ["Registrar o lead.", "Confirmar o modal."],
            "campos": [],
        }
    ]

    _, tables = _docx_estrutura(pop_minimo)
    passos = next(t for t in tables if len(t.columns) == 3 and _texto_tabela(t) == "#")
    assert [c.text for c in passos.rows[0].cells] == ["#", "Responsável", "Passo"]
    assert passos.rows[1].cells[1].text == "Comercial"
    assert passos.rows[1].cells[2].text == "Registrar o lead."
    assert passos.rows[2].cells[1].text == "Logística"


def test_banner_nao_repetido_quando_ha_responsavel_por_passo(pop_minimo: PopData) -> None:
    pop_minimo.secoes = [
        {
            "titulo": "Etapa Única",
            "responsavel": "COMERCIAL",
            "responsaveis": ["Comercial"],
            "passos": ["Agir."],
            "campos": [],
        }
    ]

    texto_tabelas = [t.rows[0].cells[0].text for t in _docx_estrutura(pop_minimo)[1]]
    assert not any("RESPONSÁVEL" in t for t in texto_tabelas)


def test_secoes_registros_criterios_indicadores_e_aviso_final(pop_minimo: PopData) -> None:
    pop_minimo.registros_obrigatorios = [
        {
            "registro": "Ficha de oportunidade",
            "conteudo": "Cliente, produto, volume",
            "responsavel": "Comercial",
        },
    ]
    pop_minimo.criterios_encerramento = (
        "Encerrar como convertida somente com primeira operação concluída."
    )
    pop_minimo.indicadores = "Taxa de conversão e tempo entre proposta e compromisso."
    pop_minimo.aviso_final = "NOTA DE CONTROLE: validar antes do uso oficial."

    texts, tables = _docx_estrutura(pop_minimo)

    assert any("Registros obrigatórios" in t for t in texts)
    assert any("Critérios de encerramento" in t for t in texts)
    assert any("Indicadores de acompanhamento" in t for t in texts)

    registros = next(t for t in tables if len(t.columns) == 3 and _texto_tabela(t) == "Registro")
    assert registros.rows[1].cells[1].text == "Cliente, produto, volume"

    aviso_final = next(t for t in tables if "NOTA DE CONTROLE" in _texto_tabela(t))
    assert "ATENÇÃO: NOTA DE CONTROLE: validar antes do uso oficial." in _texto_tabela(aviso_final)


def test_header_so_a_partir_da_segunda_pagina(pop_minimo: PopData) -> None:
    doc = Document(io.BytesIO(gerar_docx(pop_minimo).getvalue()))
    section = doc.sections[0]

    assert section.different_first_page_header_footer
    assert not section.first_page_header.paragraphs[0].text.strip()

    header_texto = section.header.paragraphs[0].text
    assert header_texto == "POP-OPE-001 · Registro de Manobras · Versão 01 · Página "
    campos = section.header.paragraphs[0]._p.findall(qn("w:fldSimple"))
    assert [c.get(qn("w:instr")) for c in campos] == ["PAGE", "NUMPAGES"]


def test_titulo_para_header() -> None:
    assert titulo_para_header("PROSPECÇÃO E FECHAMENTO COMERCIAL DE NOVAS CARGAS") == (
        "Prospecção e Fechamento Comercial de Novas Cargas"
    )
    assert titulo_para_header("PROGRAMAÇÃO DE SAÍDA") == "Programação de Saída"
    assert titulo_para_header("ATENDIMENTO À FRONTEIRA") == "Atendimento à Fronteira"
