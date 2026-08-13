"""Testes E2E do app Streamlit usando AppTest (streamlit.testing.v1).

Executa o fluxo real da interface: preencher formulario, clicar em gerar e
verificar o `.docx` produzido (mesmo caminho servido pelo botao de download).
"""

from io import BytesIO

from conftest import (
    APP_PATH,
    _abrir_formulario,
    _button,
    _download_json_historico,
    _download_pdf_gerado,
    _download_pop_atual,
    _fill_obrigatorios,
    _gerar_pop,
)
from docx import Document
from streamlit.testing.v1 import AppTest

from gerapop.constants import SessionKey, ValidationMessage
from gerapop.models import PopData
from gerapop.services.docx import gerar_docx
from gerapop.storage import get_docx_bytes, list_pops, serialize_pop
from gerapop.ui import home


def _docx_texts(pop: PopData) -> list[str]:
    doc = Document(BytesIO(gerar_docx(pop).getvalue()))
    texts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            texts.extend(cell.text for cell in row.cells)
    return texts


def test_e2e_fluxo_completo_gera_docx() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _fill_obrigatorios(at)
    _gerar_pop(at)

    assert not at.exception
    assert [success.value for success in at.success] == ["POP gerado com sucesso."]

    download = _download_pop_atual(at)
    assert len(download) == 1
    assert download[0].proto.label == "Baixar POP (.docx)"

    pop = at.session_state[SessionKey.GENERATED_POP]
    assert pop.validate() == []
    # slug truncado em 30 chars
    assert pop.output_filename() == "POP-OPE-001_Registro_de_Manobras_no_Sistem.docx"

    texts = _docx_texts(pop)
    assert "Registro de Manobras no Sistema TOS" in texts
    assert "POP-OPE-001" in texts
    assert "Operações Portuárias" in texts
    assert "Padronizar o registro de manobras." in texts


def test_e2e_geracao_com_whitespace_nas_pontas_nao_descarta_pop() -> None:
    """Regressão: campos com \n/espaços nas pontas (comum em text_area)
    eram normalizados por PopData.from_form (.strip) e a comparação
    pós-geração descartava o POP gerado (botão 'Gerar POP' sem efeito)."""
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _fill_obrigatorios(at)
    at.text_area(key="escopo").set_value("Aplica-se à equipe de operações.\n\n")
    at.text_input(key="nome_pop").set_value("Registro de Manobras \n")

    _gerar_pop(at)

    assert not at.exception
    assert [success.value for success in at.success] == ["POP gerado com sucesso."]
    assert len(_download_pop_atual(at)) == 1

    pop = at.session_state[SessionKey.GENERATED_POP]
    assert pop.escopo == "Aplica-se à equipe de operações."
    assert pop.nome_pop == "Registro de Manobras"


def test_e2e_validacao_impede_geracao() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _gerar_pop(at)

    assert not at.exception
    assert not at.success
    assert at.get("download_button") == []

    errors = [error.value for error in at.error]
    assert len(errors) == 4
    assert ValidationMessage.NOME_OBRIGATORIO in errors
    assert ValidationMessage.CODIGO_OBRIGATORIO in errors
    assert ValidationMessage.AREA_OBRIGATORIA in errors
    assert ValidationMessage.OBJETIVO_OBRIGATORIO in errors

    assert SessionKey.GENERATED_POP not in at.session_state


def test_e2e_listas_dinamicas_integradas() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _button(at, "+ Adicionar termo").click()
    at.run()
    at.text_input(key="termo_1").set_value("TOS")
    at.text_input(key="def_1").set_value("Terminal Operating System")

    at.text_input(key="sec_titulo_0").set_value("Atracação")
    at.text_input(key="passo_0_0").set_value("Verificar condições.")
    _button(at, "+ Adicionar passo").click()
    at.run()
    at.text_input(key="passo_0_1").set_value("Registrar no sistema.")

    _button(at, "+ Adicionar regra").click()
    at.run()
    at.text_input(key="regra_1").set_value("Não executar sem autorização.")

    _button(at, "+ Adicionar revisão").click()
    at.run()
    at.text_input(key="rev_1").set_value("02")
    at.text_input(key="revdesc_1").set_value("Revisão geral")

    _fill_obrigatorios(at)
    _gerar_pop(at)

    assert not at.exception
    assert [success.value for success in at.success] == ["POP gerado com sucesso."]

    pop = at.session_state[SessionKey.GENERATED_POP]
    texts = _docx_texts(pop)
    assert "TOS" in texts
    assert "Terminal Operating System" in texts
    assert "4.  Atracação" in texts
    assert "Verificar condições." in texts
    assert "Registrar no sistema." in texts
    assert "Não executar sem autorização." in texts
    assert "Revisão geral" in texts


def test_e2e_edicao_invalida_download() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _fill_obrigatorios(at)
    _gerar_pop(at)
    assert _download_pop_atual(at)

    at.text_area[0].set_value("Objetivo alterado após geração")
    at.run()

    assert _download_pop_atual(at) == []
    assert not at.success
    assert SessionKey.GENERATED_POP not in at.session_state


def test_e2e_pop_snapshot_isolado() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _fill_obrigatorios(at)
    _gerar_pop(at)
    assert _download_pop_atual(at)

    _button(at, "+ Adicionar termo").click()
    at.run()

    assert _download_pop_atual(at) == []


def test_e2e_docx_gerado_uma_vez() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _fill_obrigatorios(at)
    _gerar_pop(at)

    docx = at.session_state[SessionKey.GENERATED_DOCX]
    assert docx.getvalue()[:2] == b"PK"
    assert len(docx.getvalue()) > 1000

    at.run()
    assert at.session_state[SessionKey.GENERATED_DOCX] is docx


def test_e2e_remover_oculto_com_um_item() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    assert [b for b in at.button if b.label == "Remover"] == []

    _button(at, "+ Adicionar termo").click()
    at.run()

    assert len([b for b in at.button if b.label == "Remover"]) == 2


def test_e2e_remover_revisao() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    assert [b for b in at.button if b.label == "Remover"] == []

    _button(at, "+ Adicionar revisão").click()
    at.run()
    assert len([b for b in at.button if b.label == "Remover"]) == 2

    at.button(key="rm_rev_0").click()
    at.run()
    at.run()  # AppTest retém widgets do run pré-rerun; este run re-renderiza limpo

    assert len(at.session_state[SessionKey.REVISOES]) == 1
    assert [b for b in at.button if b.label == "Remover"] == []


def test_e2e_backup_zip_no_historico() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _fill_obrigatorios(at)
    _gerar_pop(at)
    at.run()

    backups = [d for d in at.get("download_button") if d.proto.label == "Baixar backup (.zip)"]
    assert len(backups) == 1
    assert backups[0].proto.url.endswith(".zip")


def test_e2e_rascunho_restaurado_em_nova_sessao() -> None:
    at1 = AppTest.from_file(APP_PATH, default_timeout=10)
    at1.run()

    _abrir_formulario(at1)
    at1.text_input[0].set_value("POP Rascunho Teste")
    at1.text_input[1].set_value("POP-RAS-001")

    _button(at1, "+ Adicionar campo").click()
    at1.run()
    at1.text_input(key="campo_0_0").set_value("Berço")
    at1.run()

    at2 = AppTest.from_file(APP_PATH, default_timeout=10)
    at2.run()

    _abrir_formulario(at2)
    assert at2.text_input[0].value == "POP Rascunho Teste"
    assert at2.text_input[1].value == "POP-RAS-001"
    assert at2.session_state[SessionKey.SECOES][0]["campos"][0]["campo"] == "Berço"


def test_e2e_campos_por_secao() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    at.text_input(key="sec_titulo_0").set_value("Atracação")

    _button(at, "+ Adicionar campo").click()
    at.run()
    at.text_input(key="campo_0_0").set_value("Berço")
    at.text_input(key="campodesc_0_0").set_value("Número do berço designado.")
    assert [b for b in at.button if b.label == "Remover"] == []

    _button(at, "+ Adicionar campo").click()
    at.run()
    assert len([b for b in at.button if b.label == "Remover"]) == 2
    at.text_input(key="campo_0_1").set_value("Prático")
    at.text_input(key="campodesc_0_1").set_value("Nome do prático.")

    _fill_obrigatorios(at)
    _gerar_pop(at)

    assert not at.exception
    pop = at.session_state[SessionKey.GENERATED_POP]
    assert pop.secoes[0]["campos"][1]["campo"] == "Prático"

    texts = _docx_texts(pop)
    assert "Campos obrigatórios – Atracação:" in texts
    assert "Berço" in texts
    assert "Nome do prático." in texts

    payload = serialize_pop(pop)
    assert payload["pop"]["secoes"][0]["campos"][0]["campo"] == "Berço"
    assert payload["pop"]["secoes"][0]["campos"][1]["descricao"] == "Nome do prático."


def test_e2e_export_pdf() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _fill_obrigatorios(at)
    _gerar_pop(at)

    assert not at.exception
    downloads = _download_pdf_gerado(at)
    assert len(downloads) == 1
    assert downloads[0].proto.type == "secondary"
    assert downloads[0].proto.url.endswith(".pdf")

    docx = _download_pop_atual(at)
    assert len(docx) == 1
    assert docx[0].proto.type == "primary"
    assert docx[0].proto.url.endswith(".docx")


def test_e2e_historico_export_pdf() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _fill_obrigatorios(at)
    _gerar_pop(at)
    at.run()  # re-render limpo após o st.rerun do try_generate

    downloads = _download_json_historico(at)
    assert len(downloads) == 1
    assert downloads[0].proto.url.endswith(".pdf")

    docx = [d for d in at.get("download_button") if d.proto.label == "Baixar .docx"]
    assert len(docx) == 1
    assert docx[0].proto.url.endswith(".docx")


def test_e2e_historico() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _fill_obrigatorios(at)
    _gerar_pop(at)
    at.run()  # re-render limpo após o st.rerun do try_generate

    records = list_pops()
    assert len(records) == 1
    assert records[0]["codigo"] == "POP-OPE-001"
    assert get_docx_bytes(records[0]["id"]) is not None

    _button(at, "Carregar para editar").click()
    at.run()
    at.run()  # re-render limpo após o st.rerun do carregamento

    assert at.text_input[0].value == "Registro de Manobras no Sistema TOS"
    assert at.session_state[SessionKey.OBJETIVO] == "Padronizar o registro de manobras."
    assert len(at.session_state[SessionKey.REVISOES]) == 1


def test_e2e_historico_excluir_remove_pop() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _fill_obrigatorios(at)
    _gerar_pop(at)
    at.run()  # re-render limpo após o st.rerun do try_generate

    records = list_pops()
    assert len(records) == 1
    pop_id = records[0]["id"]

    _button(at, "Excluir").click()
    at.run()
    assert "não pode ser desfeita" in at.warning[0].value

    _button(at, "Sim, excluir").click()
    at.run()
    at.run()  # re-render limpo após o st.rerun da exclusão

    assert list_pops() == []
    assert get_docx_bytes(pop_id) is None
    assert not at.exception


def test_e2e_home_excluir_pop_salvo() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _fill_obrigatorios(at)
    _gerar_pop(at)
    at.run()  # re-render limpo após o st.rerun do try_generate

    records = list_pops()
    assert len(records) == 1
    pop_id = records[0]["id"]

    at.sidebar.radio[0].set_value(home.PAGINA_HOME)
    at.run()
    assert at.button(key=f"excluir_salvo_{pop_id}") is not None

    at.button(key=f"excluir_salvo_{pop_id}").click()
    at.run()
    assert "não pode ser desfeita" in at.warning[0].value

    _button(at, "Sim, excluir").click()
    at.run()
    at.run()  # re-render limpo após o st.rerun da exclusão

    assert list_pops() == []
    assert not at.exception


def test_e2e_historico_visualizar_abre_preview() -> None:
    at = AppTest.from_file(APP_PATH, default_timeout=10)
    at.run()

    _abrir_formulario(at)
    _fill_obrigatorios(at)
    _gerar_pop(at)
    at.run()  # re-render limpo após o st.rerun do try_generate

    at.button(key="historico_ver").click()
    at.run()

    assert not at.exception
    assert at.sidebar.radio[0].value == home.PAGINA_HOME
    assert any("POP-OPE-001" in el.value for el in at.markdown)
    assert at.button(key="preview_voltar") is not None
    assert at.button(key="preview_editar") is not None
