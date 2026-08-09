"""
GeraPOP - CODEBA
MVP local single-use: preenche um formulário e gera o .docx formatado.
Rodar com: streamlit run app.py
"""
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import date

st.set_page_config(page_title="GeraPOP - CODEBA", layout="centered")

# ---------------------------------------------------------------------------
# Helpers de estado (listas dinâmicas)
# ---------------------------------------------------------------------------

def init_state():
    defaults = {
        "definicoes": [{"termo": "", "definicao": ""}],
        "secoes": [{"titulo": "", "passos": [""]}],
        "regras": [""],
        "revisoes": [{"revisao": "01", "data": str(date.today().strftime("%d/%m/%Y")),
                       "descricao": "Emissão inicial", "responsavel": ""}],
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()


def add_item(key, template):
    st.session_state[key].append(template)


def remove_item(key, idx):
    if len(st.session_state[key]) > 1:
        st.session_state[key].pop(idx)


def add_passo(sec_idx):
    st.session_state["secoes"][sec_idx]["passos"].append("")


def remove_passo(sec_idx, passo_idx):
    passos = st.session_state["secoes"][sec_idx]["passos"]
    if len(passos) > 1:
        passos.pop(passo_idx)


# ---------------------------------------------------------------------------
# Geração do .docx
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def style_header_cell(cell, text, bold=True, size=10, shading=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if shading:
        set_cell_shading(cell, shading)


def gerar_docx(dados):
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("POP – Procedimento Operacional Padrão")
    r.bold = True
    r.font.size = Pt(14)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitulo.add_run(dados["nome_pop"])
    r2.bold = True
    r2.font.size = Pt(12)

    doc.add_paragraph()

    # Tabela de cabeçalho (Código / Versão / Data / Área)
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = [("Código", dados["codigo"]), ("Versão", dados["versao"]),
              ("Data", dados["data"]), ("Área", dados["area"])]
    style_header_cell(tbl.rows[0].cells[0], "Código", shading="D9D9D9")
    style_header_cell(tbl.rows[0].cells[2], "Versão", shading="D9D9D9")
    style_header_cell(tbl.rows[1].cells[0], "Data", shading="D9D9D9")
    style_header_cell(tbl.rows[1].cells[2], "Área", shading="D9D9D9")
    tbl.rows[0].cells[1].text = dados["codigo"]
    tbl.rows[0].cells[3].text = dados["versao"]
    tbl.rows[1].cells[1].text = dados["data"]
    tbl.rows[1].cells[3].text = dados["area"]

    doc.add_paragraph()

    # Aviso (opcional)
    if dados["aviso"].strip():
        aviso_tbl = doc.add_table(rows=1, cols=1)
        aviso_tbl.style = "Table Grid"
        cell = aviso_tbl.rows[0].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run("⚠ ATENÇÃO: " + dados["aviso"])
        run.bold = True
        set_cell_shading(cell, "FFF2CC")
        doc.add_paragraph()

    def heading(numero, texto):
        h = doc.add_heading(level=1)
        h.add_run(f"{numero}.  {texto}")

    # 1. Objetivo
    heading(1, "Objetivo")
    doc.add_paragraph(dados["objetivo"])

    # 2. Escopo
    heading(2, "Escopo e Pré-condições")
    doc.add_paragraph(dados["escopo"])

    # 3. Definições
    if any(d["termo"].strip() for d in dados["definicoes"]):
        heading(3, "Definições")
        dt = doc.add_table(rows=1, cols=2)
        dt.style = "Table Grid"
        style_header_cell(dt.rows[0].cells[0], "Termo", shading="D9D9D9")
        style_header_cell(dt.rows[0].cells[1], "Definição", shading="D9D9D9")
        for item in dados["definicoes"]:
            if item["termo"].strip():
                row = dt.add_row()
                row.cells[0].text = item["termo"]
                row.cells[1].text = item["definicao"]

    # 4. Procedimento (seções dinâmicas)
    heading(4, "Procedimento")
    for i, sec in enumerate(dados["secoes"], start=1):
        if not sec["titulo"].strip():
            continue
        sub = doc.add_heading(level=2)
        sub.add_run(f"4.{i}.  {sec['titulo']}")
        pt = doc.add_table(rows=0, cols=2)
        pt.style = "Table Grid"
        pt.columns[0].width = Cm(1)
        for j, passo in enumerate(sec["passos"], start=1):
            if passo.strip():
                row = pt.add_row()
                row.cells[0].text = str(j)
                row.cells[1].text = passo

    # 5. Regras e Restrições
    if any(r.strip() for r in dados["regras"]):
        heading(5, "Regras e Restrições")
        for regra in dados["regras"]:
            if regra.strip():
                doc.add_paragraph(regra, style="List Bullet")

    # 6. Consulta e Relatórios
    if dados["consulta"].strip():
        heading(6, "Consulta e Relatórios")
        doc.add_paragraph(dados["consulta"])

    # 7. Histórico de Revisões
    heading(7, "Histórico de Revisões")
    rt = doc.add_table(rows=1, cols=4)
    rt.style = "Table Grid"
    for idx, txt in enumerate(["Revisão", "Data", "Descrição", "Responsável"]):
        style_header_cell(rt.rows[0].cells[idx], txt, shading="D9D9D9")
    for rev in dados["revisoes"]:
        if rev["revisao"].strip():
            row = rt.add_row()
            row.cells[0].text = rev["revisao"]
            row.cells[1].text = rev["data"]
            row.cells[2].text = rev["descricao"]
            row.cells[3].text = rev["responsavel"]

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# UI
# ---------------------------------------------------------------------------

st.title("📋 GeraPOP — CODEBA")
st.caption("MVP local — preencha os campos e gere o documento formatado.")

st.header("Identificação")
col1, col2 = st.columns(2)
with col1:
    nome_pop = st.text_input("Nome do POP", placeholder="Registro de Manobras no Sistema TOS – OpenPort")
    codigo = st.text_input("Código", placeholder="POP-OPE-XXX")
    versao = st.text_input("Versão", value="01")
with col2:
    area = st.text_input("Área", placeholder="Operações Portuárias")
    data_pop = st.text_input("Data", value=date.today().strftime("%d/%m/%Y"))

aviso = st.text_input("Aviso / Atenção (opcional)", placeholder="Ex: Este POP não contempla...")

st.header("1. Objetivo")
objetivo = st.text_area("Descreva o objetivo do procedimento", height=100)

st.header("2. Escopo e Pré-condições")
escopo = st.text_area("A quem se aplica / condições prévias", height=100)

st.header("3. Definições")
for i, item in enumerate(st.session_state["definicoes"]):
    c1, c2, c3 = st.columns([2, 4, 1])
    item["termo"] = c1.text_input("Termo", value=item["termo"], key=f"termo_{i}", label_visibility="collapsed", placeholder="Termo")
    item["definicao"] = c2.text_input("Definição", value=item["definicao"], key=f"def_{i}", label_visibility="collapsed", placeholder="Definição")
    if c3.button("🗑", key=f"rm_def_{i}"):
        remove_item("definicoes", i)
        st.rerun()
st.button("+ Adicionar termo", on_click=add_item, args=("definicoes", {"termo": "", "definicao": ""}))

st.header("4. Procedimento")
for si, sec in enumerate(st.session_state["secoes"]):
    st.subheader(f"Seção {si + 1}")
    sec["titulo"] = st.text_input("Título da seção", value=sec["titulo"], key=f"sec_titulo_{si}",
                                   placeholder="Ex: Procedimento – Atracação")
    for pi, passo in enumerate(sec["passos"]):
        c1, c2 = st.columns([6, 1])
        sec["passos"][pi] = c1.text_input(f"Passo {pi + 1}", value=passo, key=f"passo_{si}_{pi}",
                                           label_visibility="collapsed", placeholder=f"Passo {pi + 1}")
        if c2.button("🗑", key=f"rm_passo_{si}_{pi}"):
            remove_passo(si, pi)
            st.rerun()
    st.button("+ Adicionar passo", key=f"add_passo_{si}", on_click=add_passo, args=(si,))
    if len(st.session_state["secoes"]) > 1:
        if st.button("🗑 Remover seção", key=f"rm_sec_{si}"):
            remove_item("secoes", si)
            st.rerun()
    st.divider()
st.button("+ Adicionar seção", on_click=add_item, args=("secoes", {"titulo": "", "passos": [""]}))

st.header("5. Regras e Restrições")
for i, regra in enumerate(st.session_state["regras"]):
    c1, c2 = st.columns([6, 1])
    st.session_state["regras"][i] = c1.text_input("Regra", value=regra, key=f"regra_{i}",
                                                    label_visibility="collapsed", placeholder="Regra")
    if c2.button("🗑", key=f"rm_regra_{i}"):
        remove_item("regras", i)
        st.rerun()
st.button("+ Adicionar regra", on_click=add_item, args=("regras", ""))

st.header("6. Consulta e Relatórios")
consulta = st.text_area("Caminho / menu para consulta (opcional)", height=70)

st.header("7. Histórico de Revisões")
for i, rev in enumerate(st.session_state["revisoes"]):
    c1, c2, c3, c4 = st.columns([1, 1.5, 4, 2])
    rev["revisao"] = c1.text_input("Rev.", value=rev["revisao"], key=f"rev_{i}", label_visibility="collapsed")
    rev["data"] = c2.text_input("Data", value=rev["data"], key=f"revdata_{i}", label_visibility="collapsed")
    rev["descricao"] = c3.text_input("Descrição", value=rev["descricao"], key=f"revdesc_{i}", label_visibility="collapsed", placeholder="Descrição")
    rev["responsavel"] = c4.text_input("Responsável", value=rev["responsavel"], key=f"revresp_{i}", label_visibility="collapsed", placeholder="Responsável")
st.button("+ Adicionar revisão", on_click=add_item, args=("revisoes", {"revisao": "", "data": "", "descricao": "", "responsavel": ""}))

st.divider()

if st.button("📄 Gerar POP (.docx)", type="primary"):
    if not nome_pop.strip() or not objetivo.strip():
        st.error("Preencha ao menos o Nome do POP e o Objetivo.")
    else:
        dados = {
            "nome_pop": nome_pop, "codigo": codigo, "versao": versao,
            "data": data_pop, "area": area, "aviso": aviso,
            "objetivo": objetivo, "escopo": escopo,
            "definicoes": st.session_state["definicoes"],
            "secoes": st.session_state["secoes"],
            "regras": st.session_state["regras"],
            "consulta": consulta,
            "revisoes": st.session_state["revisoes"],
        }
        buf = gerar_docx(dados)
        st.success("POP gerado com sucesso.")
        st.download_button(
            "⬇ Baixar POP (.docx)",
            data=buf,
            file_name=f"{codigo or 'POP'}_{nome_pop[:30].replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
