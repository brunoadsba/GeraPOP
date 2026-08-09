from __future__ import annotations

import io

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from gerapop.models import PopData


def _set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _style_header_cell(
    cell,
    text: str,
    *,
    bold: bool = True,
    size: int = 10,
    shading: str | None = None,
    color: tuple[int, int, int] | None = None,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if shading:
        _set_cell_shading(cell, shading)


def gerar_docx(dados: PopData) -> io.BytesIO:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run("POP – Procedimento Operacional Padrão")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.add_run(dados.nome_pop)
    run_sub.bold = True
    run_sub.font.size = Pt(12)

    doc.add_paragraph()

    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_header_cell(tbl.rows[0].cells[0], "Código", shading="D9D9D9")
    _style_header_cell(tbl.rows[0].cells[2], "Versão", shading="D9D9D9")
    _style_header_cell(tbl.rows[1].cells[0], "Data", shading="D9D9D9")
    _style_header_cell(tbl.rows[1].cells[2], "Área", shading="D9D9D9")
    tbl.rows[0].cells[1].text = dados.codigo
    tbl.rows[0].cells[3].text = dados.versao
    tbl.rows[1].cells[1].text = dados.data
    tbl.rows[1].cells[3].text = dados.area

    doc.add_paragraph()

    if dados.aviso:
        aviso_tbl = doc.add_table(rows=1, cols=1)
        aviso_tbl.style = "Table Grid"
        cell = aviso_tbl.rows[0].cells[0]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run("ATENÇÃO: " + dados.aviso)
        run.bold = True
        _set_cell_shading(cell, "FFF2CC")
        doc.add_paragraph()

    def heading(numero: int, texto: str) -> None:
        h = doc.add_heading(level=1)
        h.add_run(f"{numero}.  {texto}")

    heading(1, "Objetivo")
    doc.add_paragraph(dados.objetivo)

    heading(2, "Escopo e Pré-condições")
    doc.add_paragraph(dados.escopo)

    if any(item["termo"].strip() for item in dados.definicoes):
        heading(3, "Definições")
        dt = doc.add_table(rows=1, cols=2)
        dt.style = "Table Grid"
        _style_header_cell(dt.rows[0].cells[0], "Termo", shading="D9D9D9")
        _style_header_cell(dt.rows[0].cells[1], "Definição", shading="D9D9D9")
        for item in dados.definicoes:
            if item["termo"].strip():
                row = dt.add_row()
                row.cells[0].text = item["termo"]
                row.cells[1].text = item["definicao"]

    heading(4, "Procedimento")
    for i, sec in enumerate(dados.secoes, start=1):
        if not sec["titulo"].strip():
            continue
        sub = doc.add_heading(level=2)
        sub.add_run(f"4.{i}.  {sec['titulo']}")
        pt = doc.add_table(rows=0, cols=2)
        pt.style = "Table Grid"
        pt.columns[0].width = Cm(1)
        for j, passo in enumerate(sec["passos"], start=1):
            if passo.strip():
                row = pt.add_row()
                row.cells[0].text = str(j)
                row.cells[1].text = passo

    if any(regra.strip() for regra in dados.regras):
        heading(5, "Regras e Restrições")
        for regra in dados.regras:
            if regra.strip():
                doc.add_paragraph(regra, style="List Bullet")

    if dados.consulta:
        heading(6, "Consulta e Relatórios")
        doc.add_paragraph(dados.consulta)

    heading(7, "Histórico de Revisões")
    rt = doc.add_table(rows=1, cols=4)
    rt.style = "Table Grid"
    for idx, txt in enumerate(["Revisão", "Data", "Descrição", "Responsável"]):
        _style_header_cell(rt.rows[0].cells[idx], txt, shading="D9D9D9")
    for rev in dados.revisoes:
        if rev["revisao"].strip():
            row = rt.add_row()
            row.cells[0].text = rev["revisao"]
            row.cells[1].text = rev["data"]
            row.cells[2].text = rev["descricao"]
            row.cells[3].text = rev["responsavel"]

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
